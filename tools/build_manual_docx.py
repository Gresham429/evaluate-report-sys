"""生成给甲方的《使用说明》docx（交付包用）。

单一可复现来源：改这里 → `uv run python tools/build_manual_docx.py` → 重出 `docs/使用说明.docx`。
用 python-docx（项目已依赖，无 pandoc）。内容对齐 v1.1.0：七类 + 委托评估协议书 + 权重可调 +
单份偏离，以及钉钉全公司同步（手机现场实勘 / 登录 / 拉取问卷 / 保存回问卷 / 审核定稿 / 共有人）。
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "使用说明.docx"
BRAND = RGBColor(0x8A, 0x1C, 0x1C)  # 正恒红，仅用于标题点缀

sys.path.insert(0, str(ROOT))
from src.version import __version__  # noqa: E402


def _set_cjk(doc: Document) -> None:
    """让正文/标题的中文走宋体、西文走 Calibri（否则默认字体渲中文偏难看）。"""
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "List Bullet", "List Number"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def steps(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def note(doc: Document, text: str) -> None:
    """灰色小字提示。"""
    para = doc.add_paragraph()
    run = para.add_run("提示： " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.size = Pt(9.5)


def table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, c in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = c
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = c


def build() -> None:
    doc = Document()
    _set_cjk(doc)

    title = doc.add_heading("", level=0)
    run = title.add_run("房地产估价报告生成系统 · 使用说明")
    run.font.color.rgb = BRAND
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    srun = sub.add_run(f"版本 v{__version__}")
    srun.italic = True
    srun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p(doc, "估价师填表（或从 Excel 导入预填）→ 选可比实例 → 系统按市场比较法重算 → "
           "产出符合公司格式的 Word 估价报告与委托评估协议书，并为每次生成留一条可复现的台账。")

    h1(doc, "一、系统概览")
    p(doc, "本系统联公司钉钉使用：估价师用钉钉小程序现场实勘，办公端扫码登录后拉取问卷、"
           "选可比实例、按市场比较法重算、出报告。估价知识（基础表）、实例库、报告编号、"
           "审核在公司范围内统一，数据汇入公司钉钉多维表、公司统一留存。")

    h1(doc, "二、启动")
    p(doc, "先把下载的压缩包完整解压到一个文件夹（不要在压缩软件里直接双击程序——"
           "那样模板没解压出来，会出不了报告），再双击程序，浏览器会自动打开操作界面。")
    p(doc, "界面有四个页签：出报告、实例库、基础表、审核。日常出报告主要用「出报告」页。")

    h1(doc, "三、开箱即用与更新")
    p(doc, "程序已内置七类估价知识（基础表：比较因素、档次描述、修正系数）与一批起步实例，"
           "首次运行自动就位——装好直接就能出报告，无需导入任何东西。")
    p(doc, "公司在钉钉多维表维护的是最新版知识。要用最新版时，到「基础表」页点「从多维表拉取」，"
           "把公司最新版拉到本机（缺则补、已有不动、旧版永不覆盖，日后仍能复现旧报告）。")
    note(doc, "内置的是分发时的快照；日常以公司多维表为准，点「从多维表拉取」更新即可。")

    h1(doc, "四、出一份报告")
    p(doc, "支持七类：农用 / 办公 / 商业 / 住宅 / 工业 / 停车场用地 / 建设用地。流程七步：")
    steps(doc, [
        "开始：新建空白表单，或导入 Excel 预填（联钉钉时还可「从实勘问卷拉取」）。",
        "基本信息：19 项，全部人工填。",
        "估价对象一览表：可增删行。",
        "因素档次：下拉选择，选项来自基础表。",
        "选可比实例：勾三条 + 逐条填市场状况指数 → 重算 → 用结果更新一览表。",
        "附件：图片或 PDF，可跳过。",
        "生成下载：可分别下载「估价报告」与「委托评估协议书」两份 Word 文书，点哪个下哪个。",
    ])
    note(doc, "住宅 / 工业 / 停车场用地 / 建设用地 四类的报告正文为结构化渲染，"
              "页脚标注「须执业估价师终审」——交付客户前请由执业估价师复核用词与结论。")

    h2(doc, "第伍步：重算的三条规矩")
    bullets(doc, [
        "系统不推算、不给默认值：市场状况指数须你逐条填，不填就报错、不会静默取值继续算。",
        "系统不做可比性推荐：哪条更可比由你的专业判断决定；列表只按租期起始日从新到旧排，不打分、不高亮。",
        "三条实例的权重可调：默认各三分之一，可改（支持分数输入如 1/3），三者之和必须等于 1，否则挡住生成。",
    ])
    p(doc, "算完点「用此结果更新一览表」，评估结果才写进一览表单价，报告印的就是这个数。")

    h2(doc, "单份系数偏离与单价可改")
    bullets(doc, [
        "某份报告要临时调某个修正系数，可在该报告里直接改（只对这一份生效，不改基础表版本），"
        "超出建议范围只提醒、不拦，可填理由留痕。",
        "一览表单价填进去后可直接改（如把 1399.26 按 1400 写进报告）——取整与否是你的判断，"
        "系统不替你决定；改完年租赁价值自动跟着算。",
    ])

    h1(doc, "五、实例库 / 基础表版本 / 台账 / 草稿")
    h2(doc, "实例库")
    bullets(doc, [
        "从 Excel 批量导入：传一份「实勘表、比较法.xlsx」，自动抽出其中三条实例。",
        "手工录入一条：填位置、成交价、面积、租期，再选各因素档次；编号、起始日由系统据租期原文派生。",
        "租期只知年月就写「2025.7-2026.7」，系统会标「日期仅年月」把事实留在库里；重复编号不覆盖已有的。",
    ])
    h2(doc, "基础表版本")
    p(doc, "基础表由公司统一维护，以内容指纹作版本号（内容一变即变，改不了也伪造不了）。"
           "更新后推到多维表、各人再拉取即取到新版；旧版本永不覆盖——日后复查旧报告时，"
           "能拿回当时那版基础表把结果重算出来。")
    h2(doc, "台账")
    p(doc, "每生成一份报告，系统自动留一条完整快照记录（何时、谁、用哪版基础表、哪三条实例、"
           "指数填了多少、算出什么数），不用你操作。同一份报告多次生成合并成一行、显示「共 N 次」。"
           "点开任一条可「照此重算」当场验证能否复现，且全程不碰实例库和基础表库。台账只增不改。")
    h2(doc, "草稿")
    p(doc, "边填边存，不用手动保存。填一半关了浏览器，下次在第壹步能看到「未完成的草稿」，点「续填」接着填。"
           "附件不进草稿，续填后须重挑。")

    h1(doc, "六、钉钉全公司协作（联钉钉时）")
    bullets(doc, [
        "手机现场实勘：估价师用钉钉小程序在现场填实勘问卷，可地图预填地理事实、拍照、离线填写联网后自动补传。",
        "选共有人：填问卷时从钉钉通讯录选「共有人」，被选的同事在其办公端也能看到、编辑这份问卷（填报人恒在、不可删）。",
        "办公端登录：办公端用钉钉扫码登录，按登录身份决定看得到哪些问卷（只看与自己相关的）。",
        "从实勘问卷拉取：办公端「从实勘问卷拉取」把现场提交的问卷预填进出报告表单；同一问卷重复拉取会续到同一份草稿，不重复新建。",
        "保存回问卷：办公端编辑后可「保存回问卷」，与线上做字段级合并；两边改了同一字段会弹出冲突框，逐项选择保留哪个。",
        "审核定稿：问卷走「草稿 → 已提交 → 待审核 → 已定稿」四态；已定稿即锁定、只读，不可再改。",
    ])
    note(doc, "手机端「我的问卷」列出你有权查看/编辑的全部问卷（含被别人加为共有人的），不只你自己建的。")

    h1(doc, "七、重要约束")
    bullets(doc, [
        "改完 Excel 一定要保存再导入：系统读的是 Excel 存储的计算结果，不保存就导入读到的是旧值（导入基础表同理）。",
        "表单里的修改只影响本次报告，不写回 Excel：Excel 是一张公式网，回写单个值会打断公式链。",
        "校验提示只是提示、不阻止生成：是不是问题、要不要改由你判断；系统替你核，但不替你改、不替你填。",
        "新四类（住宅/工业/停车场用地/建设用地）报告须执业估价师终审用词与结论后再交付。",
    ])

    h1(doc, "八、数据与措辞")
    h2(doc, "数据放在哪")
    p(doc, "问卷、实例、台账、报告编号汇总在公司钉钉多维表，公司统一留存、统一编号；"
           "基础表由公司维护、各人拉取到本机缓存使用。本机不留敏感汇总，换机重新登录 + 拉取即可。")
    h2(doc, "改措辞")
    table(doc, ["改什么", "改哪里"], [
        ["法定套话（声明、假设限制条件、估价原则等三类通用文字）", "改 copy.yaml，改一处三类同步"],
        ["类别专属内容与格式", "用 Word 改 templates 文件夹里对应类别的模板"],
        ["因素、档次、修正系数", "由公司维护基础表 Excel、更新后各人重新「从多维表拉取」"],
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"已生成：{OUT}（{OUT.stat().st_size} 字节）")


if __name__ == "__main__":
    build()
