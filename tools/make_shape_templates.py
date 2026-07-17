"""从现有已 jinja 化模板派生两个通用 shape 模板，替换残留硬编码样例值。

office.docx→lease_building.docx（房产 shape）、farmland.docx→lease_land.docx（土地 shape）。

现有三份模板是从各自真实报告手改的，仍有几处样例值没模板化（价值时点日期、
委托人住址/法人、土地模板里的现场查勘人姓名）——直接复用会把某一个样例的
值烙进新类别报告。本脚本把这些残留值替换成 build_context 已提供的
`{{ 变量 }}`，使新 shape 模板真正通用。纯脚本、可重复运行（每次从源模板重派生）。

⚠️ 新 shape 模板正文是执业文书，Approach A 下只做结构渲染、须执业估价师终审，
故末尾追加一行终审提示（tests/test_render_new_categories.py 钉死其存在）。
"""

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType

TEMPLATES = Path(__file__).resolve().parents[1] / "templates"

REVIEW_NOTICE = (
    "本报告由系统按结构渲染生成，未经逐段金样校验；交付前须执业估价师终审用词与格式。"
)

# 精确原文 → jinja token。build_context 均已提供对应变量。
BUILDING_REPLACE = {
    "2026年3月26日": "{{ value_date_cn }}",
    "浙江省杭州市萧山区金城路685号综合楼三楼": "{{ client_address }}",
    "余峰": "{{ legal_rep }}",
}
LAND_REPLACE = {
    "2026年4月20日": "{{ value_date_cn }}",
    "杭州市钱塘区义蓬街道义蓬村": "{{ client_address }}",
    "周国祥": "{{ legal_rep }}",
    "郑伟娜": "{{ surveyor }}",
}


def _fix_paragraph(paragraph: object, mapping: dict[str, str]) -> None:
    """段落内整串替换。跨 run 的串先并到首 run 再替换（保留段落、丢弃分段格式）。"""
    text = paragraph.text  # type: ignore[attr-defined]
    if not any(k in text for k in mapping):
        return
    for old, new in mapping.items():
        text = text.replace(old, new)
    runs = paragraph.runs  # type: ignore[attr-defined]
    for run in runs:
        run.text = ""
    if runs:
        runs[0].text = text
    else:
        paragraph.add_run(text)  # type: ignore[attr-defined]


def _replace_everywhere(doc: DocumentType, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        _fix_paragraph(paragraph, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fix_paragraph(paragraph, mapping)


def _restore_styles(src: Path, dst: Path) -> None:
    """把源模板的 word/styles.xml 原样塞回派生模板。

    python-docx 保存时会重排 styles.xml 的字节（内容等价、字节不同），使
    `test_templates_share_styles` 的「全模板共用同一样式表」摔掉。派生只改了
    正文文字、没加新样式，故把源的 styles.xml 原样覆盖回去，样式字节保持一致。
    """
    styles = zipfile.Path(src, "word/styles.xml").read_bytes()
    tmp = dst.with_suffix(".tmp.docx")
    with zipfile.ZipFile(dst) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = styles if item.filename == "word/styles.xml" else zin.read(item.filename)
            zout.writestr(item, data)
    tmp.replace(dst)


def build(src_name: str, dst_name: str, mapping: dict[str, str]) -> None:
    src = TEMPLATES / src_name
    dst = TEMPLATES / dst_name
    shutil.copy(src, dst)
    doc = Document(str(dst))
    _replace_everywhere(doc, mapping)
    doc.add_paragraph(REVIEW_NOTICE)
    doc.save(str(dst))
    _restore_styles(src, dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    build("office.docx", "lease_building.docx", BUILDING_REPLACE)
    build("farmland.docx", "lease_land.docx", LAND_REPLACE)
