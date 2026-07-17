"""新类别报告结构渲染：能渲染、关键槽已填、无样例值泄漏、无未渲染 jinja、有终审标注。

Approach A：新类别不做正文金样（无签发报告可比），只保证结构正确 + 交付标注存在。
正文用词的执业合规性由估价师终审兜底（见 docs/README §8）。
"""

from pathlib import Path

import docx
import pytest

from src.engine.annual import annual_value
from src.model import Category, Project, Subject

NEW = [
    Category.RESIDENTIAL,
    Category.INDUSTRIAL,
    Category.PARKING_LAND,
    Category.CONSTRUCTION_LAND,
]
# 派生模板必须去掉的源模板样例值（人名/住址/日期）。
LEAK = ["余峰", "周国祥", "金城路685号", "义蓬街道义蓬村", "郑伟娜", "2026年3月26日", "2026年4月20日"]


def _proj(category: Category) -> Project:
    area = 342.59
    price = 1.2
    return Project(
        category=category,
        report_no="正恒评报字[2026]第TEST号",
        project_name="测试项目",
        client="某某委托人",
        client_address="某某住址路1号",
        legal_rep="某某法人",
        purpose="评估房地产租赁价值",
        survey_date="2026-05-01",
        value_date="2026-05-01",
        materials="《不动产权证》",
        certificate_status="估价对象已取得《不动产权证》",
        owner="某某权利人",
        address="某某坐落",
        usage="测试用途",
        scale="面积342.59",
        scope="包含使用权",
        current_status="至价值时点，估价对象处于空置状态。",
        work_period="2026年5月",
        issue_date="2026-05-10",
        surveyor="某某勘查人",
        unit_price=price,
        dispersion=1.0,
        subjects=(
            Subject(
                index=1,
                owner="某某权利人",
                address="某某坐落",
                usage="测试",
                area=area,
                unit_price=price,
                annual_value=annual_value(category, area, price),
            ),
        ),
    )


def _render_text(category: Category, tmp_path: Path) -> str:
    # 延迟导入，确保读的是 wire 后的 TEMPLATE_FILENAMES。
    from src.renderer.render import render

    out = render(_proj(category), (), tmp_path / "r.docx")
    assert out.exists()
    doc = docx.Document(out)
    parts = [p.text for p in doc.paragraphs]
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@pytest.mark.parametrize("category", NEW)
def test_new_category_renders_key_slots(category: Category, tmp_path: Path) -> None:
    text = _render_text(category, tmp_path)
    assert "某某权利人" in text  # owner 已填
    assert "某某坐落" in text  # address 已填
    assert "某某法人" in text  # legal_rep 已填（派生模板已模板化）


@pytest.mark.parametrize("category", NEW)
def test_new_category_no_unrendered_jinja(category: Category, tmp_path: Path) -> None:
    text = _render_text(category, tmp_path)
    assert "{{" not in text and "{%" not in text


@pytest.mark.parametrize("category", NEW)
def test_new_category_no_sample_leak(category: Category, tmp_path: Path) -> None:
    text = _render_text(category, tmp_path)
    for leaked in LEAK:
        assert leaked not in text, f"派生模板残留样例值：{leaked}"


@pytest.mark.parametrize("category", NEW)
def test_new_category_has_review_notice(category: Category, tmp_path: Path) -> None:
    text = _render_text(category, tmp_path)
    assert "须执业估价师终审" in text
