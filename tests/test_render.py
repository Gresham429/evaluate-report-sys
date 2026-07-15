"""渲染器测试。"""

import html
import re
import zipfile
from pathlib import Path

import pytest

from src.attachments.collector import collect
from src.extractor.project import load_project
from src.model import Project
from src.renderer.render import build_context, render
from tests.conftest import CASES, MATERIALS


def extract_paragraphs_for_test(path: Path) -> list[str]:
    """抽取 docx 的全部可见文本（含表格单元格），逐段返回。

    与 tools/extract_copy.py 的 paragraphs() 同一思路：<w:p> 段落既出现在
    正文流也出现在表格单元格内，把 </w:p> 归一为换行、再剥标签，表格内
    文字自然一并被抽到——不能用只遍历 document.paragraphs 的写法，那会
    漏掉表格。
    """
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    xml = re.sub(r"</w:p>", "\n", xml)
    text = html.unescape(re.sub(r"<[^>]+>", "", xml))
    return [line.strip() for line in text.split("\n") if line.strip()]


def test_context_has_subjects() -> None:
    project = load_project(CASES["办公"])
    context = build_context(project, [])
    assert len(context["subjects"]) == 2
    assert context["report_no"] == "正恒评报字[2026]第F071号"
    assert context["面积单位"] == "㎡"


def test_context_agricultural_units() -> None:
    context = build_context(load_project(CASES["农用"]), [])
    assert context["面积单位"] == "亩"
    assert context["单价单位"] == "元/亩·年"


def test_context_has_attachments_flag() -> None:
    assert build_context(load_project(CASES["办公"]), [])["has_attachments"] is False


@pytest.mark.parametrize("case", ["农用", "办公", "商业"])
def test_render_produces_valid_docx(case: str, tmp_path: Path) -> None:
    project = load_project(CASES[case])
    output = tmp_path / f"{case}.docx"
    render(project, [], output)
    assert output.exists()
    with zipfile.ZipFile(output) as archive:
        assert "word/document.xml" in archive.namelist()


def test_render_is_deterministic(tmp_path: Path) -> None:
    """同一输入渲染两次，正文必须逐字节相同（约束 C2）。"""
    project = load_project(CASES["商业"])
    first, second = tmp_path / "a.docx", tmp_path / "b.docx"
    render(project, [], first)
    render(project, [], second)
    with zipfile.ZipFile(first) as fa, zipfile.ZipFile(second) as fb:
        assert fa.read("word/document.xml") == fb.read("word/document.xml")


def test_render_with_attachments(tmp_path: Path) -> None:
    project = load_project(CASES["办公"])
    pages = collect([MATERIALS / "办公" / "附件" / "海创703.pdf"], tmp_path / "img")
    output = tmp_path / "with.docx"
    render(project, pages, output)
    with zipfile.ZipFile(output) as archive:
        media = [n for n in archive.namelist() if n.startswith("word/media/")]
    assert len(media) >= 7


def _fake_office_project() -> Project:
    """构造一个用伪造数据替换办公金样的 Project。

    address 也一并伪造（原始报告里的 test 草稿没有伪造它）——project.address
    是与 subject.address 不同的字段（前者是整份报告的项目地址，后者是每个
    估价对象的门牌号），二者都已被正确参数化。若不伪造 project.address，
    真实地址会经由正确工作的 {{ address }} 合法出现在文本里，跟"模板里
    写死了金样数据"这个判定会混在一起，分不清是真泄漏还是没伪造到位。
    """
    from dataclasses import replace

    from src.model import Subject

    project = load_project(CASES["办公"])
    return replace(
        project,
        owner="测试公司XYZ",
        report_no="正恒评报字[9999]第TEST号",
        address="测试地址999号",
        subjects=(
            Subject(
                index=1, owner="测试公司XYZ", address="测试地址999室", usage="办公",
                area=111.11, unit_price=9.99, annual_value=405147,
            ),
        ),
    )


def test_template_has_no_hardcoded_golden_data(tmp_path: Path) -> None:
    """用伪造数据渲染，金样的数字必须全部消失。

    这是唯一能发现「模板里写死了金样数据」的测试——所有拿金样自己的项目
    去渲染再跟金样比的测试（含金样回归）都验不出这类缺陷。

    覆盖范围：Task 12 修复的「估价结果一览表」及两张摘要表的数据行，以及
    随手一并修好的封面标题地址（同一 bug 类：SUBSTITUTIONS 原先对整份
    document.xml 做一次性字符串替换，被 WPS 插入的空 <w:bookmarkStart>/
    <w:bookmarkEnd> 打断成跨 run 短语的地方会静默替换失败）。

    不在本次覆盖范围内、已确认另外存在、且仍会泄漏金样数字的 3 处，见
    test_known_out_of_scope_golden_leaks()与 task-12-fix-report.md：
    「实物状况/建筑规模」说明性文字表格（按类别措辞结构不同，需要新的
    条件文案设计）、"估价范围"/"依据不足假设"两段里从未接上 {{ scale }}
    占位符的面积描述、"年租赁价值为...元，大写：人民币...元整"一句里
    同时硬编码的阿拉伯数字与中文大写金额（后者需要新增人民币大写转换
    能力，本次改动未涉及）。
    """
    fake = _fake_office_project()
    output = tmp_path / "fake.docx"
    render(fake, [], output)
    text = "\n".join(extract_paragraphs_for_test(output))

    # 伪造数据必须出现
    assert "测试公司XYZ" in text
    assert "测试地址999号" in text
    assert "测试地址999室" in text
    assert "111.11" in text
    assert "9.99" in text
    assert "405,147" in text

    # 金样数据必须彻底消失（一览表/摘要表数据行 + 封面标题地址）
    for leaked in ("杭州萧山国有资产投资有限公司", "368030", "368,030",
                   "379506", "379,506", "萧山区北干街道萧山科创中心3幢1206室",
                   "萧山科创中心3幢1206室"):
        assert leaked not in text, f"模板里残留金样数据：{leaked}"


@pytest.mark.xfail(strict=True, reason="Task 12 范围外的既存缺陷，见 task-12-fix-report.md")
def test_known_out_of_scope_golden_leaks(tmp_path: Path) -> None:
    """记录 Task 12 验证过程中新发现、但明确不在本次修复范围内的金样泄漏。

    用 xfail(strict=True) 而不是直接删除断言或默默放过：这些是真实存在的
    「模板写死金样数据」缺陷，与 Task 12 同一 bug 类，但都需要新的条件
    文案设计或新的中文大写金额转换能力，涉及未经用户确认的措辞/生成逻辑
    决策，不应在本次改动中顺手臆造。保留为 xfail 让它们在测试套件里可见、
    可追踪——一旦哪天被修好，这个测试会变成"意外通过"而失败，提醒把它
    转成正式断言，而不是被无声遗忘。
    """
    fake = _fake_office_project()
    output = tmp_path / "fake.docx"
    render(fake, [], output)
    text = "\n".join(extract_paragraphs_for_test(output))
    for leaked in (
        "356.29",  # 实物状况/建筑规模 表格：分室面积明细，按类别措辞结构不同
        "367.4",
        "723.69",  # 同上表格 + "估价范围"/"依据不足假设"两段（scale 从未接上占位符）
        "萧山科创中心3幢1208室",  # 同上表格
        "747536",  # "年租赁价值为...元，大写：..."句：阿拉伯数字+中文大写金额均硬编码
    ):
        assert leaked not in text, f"模板里残留金样数据：{leaked}"


def test_subject_count_drives_table_rows(tmp_path: Path) -> None:
    """农用 1 个对象、办公 2 个——表格行数必须随数据变，不能写死。"""
    from docx import Document

    for case, expected in (("农用", 1), ("办公", 2)):
        project = load_project(CASES[case])
        assert len(project.subjects) == expected
        output = tmp_path / f"{case}.docx"
        render(project, [], output)
        document = Document(str(output))
        summary = [
            t for t in document.tables
            if any("年租赁价值" in c.text for c in t.rows[0].cells)
        ]
        assert summary, f"{case} 未找到估价结果一览表"
        # 表头 + expected 个数据行 + 合计行
        assert len(summary[0].rows) == expected + 2, (
            f"{case} 一览表应有 {expected + 2} 行，实为 {len(summary[0].rows)}"
        )
